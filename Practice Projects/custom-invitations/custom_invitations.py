#! python3
# custom_invitations.py - Generates a Word document with custom invitations


import docx


guests = open('guests.txt').readlines()
doc = docx.Document()

for guest in guests:

    doc.add_paragraph('It would be a pleasure to have the company of', style='Quote')
    doc.add_paragraph(guest.strip().upper())
    doc.add_paragraph('at 11010 Memory Lane on Evening of', style='Quote')
    doc.add_paragraph('April 1st')
    doc.add_paragraph('at 7 o’clock', style='Quote')

    if guest != guests[-1]:
        doc.add_page_break()

doc.save('invitations.docx')
